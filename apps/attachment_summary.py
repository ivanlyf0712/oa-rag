#!/usr/bin/env python3
"""Per-contract summary that reads the signed attachment file (PDF/DOCX)
and combines it with the recorded risk tags into a two-section review summary.

Design (grilled decisions):
- Attachment chosen by field priority: signedcontract -> finalversioncontract -> DraftContract.
- PDF extracted via PyMuPDF (fitz); DOCX via python-docx.
- Text capped (~10k chars) with a truncation note.
- Graceful fallback: unreadable/missing/scanned file -> risk-tags-only summary + notice.
- Result cached per (ref_no, file mtime) so repeat views are instant.
- Attachment lookup joins contract_id -> main id, falls back to ref_no (logged).
"""

from __future__ import annotations

import json
import logging
import os
from typing import Any, Callable, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# Path resolution (host-absolute DB paths vs container /app/uploads)
_UPLOADS_MARKER = "uploads/"


def _default_uploads_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "uploads"))


def uploads_root() -> str:
    """Configured uploads root: UPLOADS_ROOT env var, else the repo uploads/ dir."""
    return os.environ.get("UPLOADS_ROOT") or _default_uploads_root()


def resolve_attachment_path(path: str, root: Optional[str] = None) -> str:
    """Return path as-is when it exists on disk; otherwise re-anchor the
    portion from the uploads/ marker onward under the configured uploads
    root. Paths without an uploads/ marker are returned unchanged."""
    p = (path or "").strip()
    if not p or os.path.isfile(p):
        return p
    norm = p.replace(chr(92), "/")
    if "/" + _UPLOADS_MARKER in norm:
        suffix = norm.split("/" + _UPLOADS_MARKER, 1)[1]
    elif norm.startswith(_UPLOADS_MARKER):
        suffix = norm[len(_UPLOADS_MARKER):]
    else:
        return p
    return os.path.join(root or uploads_root(), *suffix.split("/"))


# Humanized attachment labels (OA field_name -> plain English)
ATTACHMENT_FIELD_LABELS: Dict[str, str] = {
    "signedcontract": "Signed contract",
    "finalversioncontract": "Final version",
    "draftcontract": "Draft",
    "unspecified": "Other attachment",
}


def attachment_label(field_name: Optional[str]) -> str:
    """Human-readable label for an OA attachment field name."""
    key = (field_name or "").strip().lower()
    if key in ATTACHMENT_FIELD_LABELS:
        return ATTACHMENT_FIELD_LABELS[key]
    return key.replace("_", " ").title() if key else "Other attachment"


def human_file_size(size: Any) -> str:
    """Human-readable file size; empty string when unknown."""
    try:
        n = float(size)
    except (TypeError, ValueError):
        return ""
    if n <= 0:
        return ""
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if n < 1024 or unit == "TB":
            return ("%d %s" % (n, unit)) if unit == "B" else ("%.1f %s" % (n, unit))
        n /= 1024
    return ""


# Attachment field priority: the signed contract is the binding source of truth.
ATTACHMENT_FIELD_PRIORITY: Tuple[str, ...] = (
    "signedcontract",
    "finalversioncontract",
    "DraftContract",
)

DEFAULT_TEXT_CAP = 10000

_PDF_MIME = "application/pdf"
_DOCX_MIMES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
)


# ─────────────────────────────────────────────────────────────────────
# Attachment lookup (DB)
# ─────────────────────────────────────────────────────────────────────
def _query_attachments(contract_id: Optional[Any], ref_no: Optional[str],
                       db_connect: Optional[Callable] = None) -> List[Dict[str, Any]]:
    """Fetch attachment rows for a contract from contract_attachments.

    Tries contract_id -> main id first, then ref_no. db_connect is injectable for tests.
    """
    if db_connect is None:
        from core.db import get_db_connection
        db_connect = get_db_connection

    conn = db_connect()
    try:
        cur = conn.cursor()
        rows: List[Dict[str, Any]] = []
        if contract_id is not None and str(contract_id).strip() != "":
            cur.execute(
                "SELECT field_name, file_name, file_path, mime_type, file_size "
                "FROM contract_attachments WHERE contract_id = %s",
                (contract_id,),
            )
            cols = [d[0] for d in cur.description]
            rows = [dict(r) if isinstance(r, dict) else dict(zip(cols, r)) for r in cur.fetchall()]
            if rows:
                logger.debug("attachment lookup matched on contract_id=%s", contract_id)
                return rows
        if ref_no:
            cur.execute(
                "SELECT field_name, file_name, file_path, mime_type, file_size "
                "FROM contract_attachments WHERE ref_no = %s",
                (ref_no,),
            )
            cols = [d[0] for d in cur.description]
            rows = [dict(r) if isinstance(r, dict) else dict(zip(cols, r)) for r in cur.fetchall()]
            if rows:
                logger.debug("attachment lookup matched on ref_no=%s", ref_no)
        return rows
    finally:
        try:
            conn.close()
        except Exception:
            pass


def list_attachments(contract_id: Optional[Any] = None, ref_no: Optional[str] = None,
                     db_connect: Optional[Callable] = None) -> List[Dict[str, Any]]:
    """Public wrapper around _query_attachments: all attachment rows for a contract.

    Returns rows with field_name, file_name, file_path, mime_type, file_size.
    """
    return _query_attachments(contract_id, ref_no, db_connect)


def choose_attachment(
    contract_id: Optional[Any],
    ref_no: Optional[str],
    field_priority: Tuple[str, ...] = ATTACHMENT_FIELD_PRIORITY,
    db_connect: Optional[Callable] = None,
) -> Optional[Dict[str, Any]]:
    """Pick the highest-priority attachment whose file exists on disk.

    Returns the attachment row dict or None if no usable file is found.
    """
    try:
        rows = _query_attachments(contract_id, ref_no, db_connect=db_connect)
    except Exception as e:
        logger.warning("attachment lookup failed (contract_id=%s ref_no=%s): %s",
                       contract_id, ref_no, e)
        return None
    if not rows:
        return None

    def _resolved(row: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        path = (row.get("file_path") or "").strip()
        if not path:
            return None
        cand_path = resolve_attachment_path(path)
        if os.path.isfile(cand_path):
            row["file_path"] = cand_path
            return row
        return None

    by_field: Dict[str, Dict[str, Any]] = {}
    for r in rows:
        field = (r.get("field_name") or "").strip().lower()
        by_field.setdefault(field, r)

    for field in field_priority:
        cand = by_field.get(field.strip().lower())
        if not cand:
            continue
        hit = _resolved(cand)
        if hit is not None:
            return hit
    # nothing matched priority with an existing file; fall back to any existing file
    for r in rows:
        hit = _resolved(r)
        if hit is not None:
            return hit
    return None


# ─────────────────────────────────────────────────────────────────────
# Text extraction (PDF via fitz, DOCX via python-docx)
# ─────────────────────────────────────────────────────────────────────
def _extract_pdf(path: str) -> str:
    import fitz  # PyMuPDF
    parts: List[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts)


def _extract_docx(path: str) -> str:
    import docx  # python-docx
    document = docx.Document(path)
    parts: List[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(
    path: str,
    mime_type: Optional[str] = None,
    cap: int = DEFAULT_TEXT_CAP,
) -> Dict[str, Any]:
    """Extract plain text from a PDF/DOCX attachment.

    Returns {"text": str, "truncated": bool, "chars": int, "error": str|None}.
    On failure, error is set and text is empty (caller decides fallback).
    """
    result: Dict[str, Any] = {"text": "", "truncated": False, "chars": 0, "error": None}
    if not path or not os.path.isfile(path):
        result["error"] = "file not found"
        return result

    mt = (mime_type or "").lower()
    ext = os.path.splitext(path)[1].lower()
    try:
        if mt == _PDF_MIME or ext == ".pdf":
            text = _extract_pdf(path)
        elif mt in _DOCX_MIMES or ext in (".docx", ".doc"):
            text = _extract_docx(path)
        else:
            # best-effort: try pdf then docx
            try:
                text = _extract_pdf(path)
            except Exception:
                text = _extract_docx(path)
    except Exception as e:
        result["error"] = "extraction failed: %s" % e
        return result

    text = (text or "").strip()
    if not text:
        result["error"] = "no extractable text (scanned image or empty file)"
        return result

    result["chars"] = len(text)
    if cap and len(text) > cap:
        result["text"] = text[:cap]
        result["truncated"] = True
    else:
        result["text"] = text
    return result


# ─────────────────────────────────────────────────────────────────────
# LLM invocation (reuses the agent's self-healing chain)
# ─────────────────────────────────────────────────────────────────────
def _call_summary_llm(llm: Any, system_msg: str, user_msg: str) -> Optional[str]:
    """Invoke the chat model with a plain (non-tool) message; return text or None."""
    if llm is None:
        return None
    try:
        resp = llm.invoke([("system", system_msg), ("human", user_msg)])
        text = getattr(resp, "content", "") or (resp if isinstance(resp, str) else "")
        text = (text or "").strip()
        return text or None
    except Exception as e:
        logger.warning("attachment summary LLM call failed: %s", e)
        return None


# ─────────────────────────────────────────────────────────────────────
# Two-section summary: document summary (from file) + risk assessment (from tags)
# ─────────────────────────────────────────────────────────────────────
def _summarize_document(text: str, meta: Dict[str, Any], llm: Any,
                        truncated: bool) -> Tuple[Optional[str], bool]:
    """LLM-summarize the extracted attachment text. Returns (summary, llm_used)."""
    if not text:
        return None, False
    name = meta.get("title") or meta.get("ref_no") or "this contract"
    system_msg = (
        "You are a legal reviewer summarizing a signed contract document. "
        "Use ONLY the document text provided. Write one concise paragraph (3-5 sentences) "
        "covering: the parties, the subject matter / scope, the contract value and term if stated, "
        "and any notable obligations, liabilities, or unusual clauses. "
        "Do not invent facts not present in the text. Plain English, no record IDs or field names."
    )
    trunc_note = " (Note: the document was truncated to its opening portion.)" if truncated else ""
    user_msg = (
        "Contract: " + str(name) + "." + trunc_note + "\n\n"
        "Document text:\n" + text + "\n\nSummary:"
    )
    out = _call_summary_llm(llm, system_msg, user_msg)
    return (out, bool(out))


def summarize_contract_with_attachment(
    row: Any,
    llm: Any = None,
    text_cap: int = DEFAULT_TEXT_CAP,
    field_priority: Tuple[str, ...] = ATTACHMENT_FIELD_PRIORITY,
    db_connect: Optional[Callable] = None,
    cache: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Produce a two-section review summary for one contract row.

    Section 1 ("Document summary") is generated from the signed attachment file.
    Section 2 ("Risk assessment") comes from the recorded risk tags (existing
    build_contract_context / deterministic / LLM logic reused).

    Falls back to a risk-tags-only summary with a notice when the attachment
    cannot be read. Results are cached per (ref_no/id, attachment mtime).

    Returns a dict with keys:
      document_summary, risk_summary, summary (combined), context, attachment
      (file_name/field_name/path or None), llm_used, notice (str or None).
    """
    # Local imports to avoid a hard pandas dependency at module import time.
    from apps.risk_search import build_contract_context, summarize_contract

    meta = row if isinstance(row, dict) else row.to_dict()
    contract_id = meta.get("id")
    ref_no = meta.get("ref_no") or meta.get("RefNo")

    # Always build the risk-tag context (cheap) so we can fall back to it.
    import pandas as pd
    row_series = row if hasattr(row, "index") else pd.Series(meta)
    context = build_contract_context(row_series)

    attachment = choose_attachment(contract_id, ref_no, field_priority=field_priority,
                                   db_connect=db_connect)

    # ── cache key: contract + attachment mtime ──────────────────────────
    cache_key = None
    if cache is not None:
        mtime = None
        if attachment and attachment.get("file_path"):
            try:
                mtime = os.path.getmtime(attachment["file_path"])
            except OSError:
                mtime = None
        cache_key = "%s|%s" % (ref_no or contract_id, mtime)
        if cache_key in cache:
            return cache[cache_key]

    notice: Optional[str] = None
    document_summary: Optional[str] = None
    doc_llm_used = False

    if attachment is None:
        notice = "No readable attachment was found for this contract, so the summary is based on the recorded risk tags only."
    else:
        extraction = extract_text(attachment.get("file_path"), attachment.get("mime_type"), cap=text_cap)
        if extraction.get("error"):
            notice = ("The attachment could not be read (%s), so the summary is based on the "
                      "recorded risk tags only." % extraction["error"])
        else:
            document_summary, doc_llm_used = _summarize_document(
                extraction["text"], context.get("metadata", {}), llm, extraction.get("truncated", False)
            )
            if document_summary is None:
                notice = "The language model is unavailable, so the document could not be summarized; the risk assessment below is based on the recorded risk tags."

    # ── risk assessment from recorded tags (reuse existing, self-contained logic) ──
    # summarize_contract runs its own RiskPlanner LLM path with a deterministic
    # fallback, so the risk section is correct whether or not a document LLM is set.
    risk_out = summarize_contract(row_series, use_llm=True)
    risk_summary = risk_out["summary"]
    risk_llm_used = bool(risk_out.get("llm_used"))

    # ── combine sections ───────────────────────────────────────────────
    parts: List[str] = []
    if document_summary:
        parts.append("**Document summary (from the signed attachment)**\n\n" + document_summary)
    parts.append("**Risk assessment (from recorded risk tags)**\n\n" + risk_summary)
    combined = "\n\n".join(parts)

    result: Dict[str, Any] = {
        "document_summary": document_summary,
        "risk_summary": risk_summary,
        "summary": combined,
        "context": context,
        "attachment": (
            {
                "file_name": attachment.get("file_name"),
                "field_name": attachment.get("field_name"),
                "file_path": attachment.get("file_path"),
            }
            if attachment else None
        ),
        "llm_used": bool(doc_llm_used or risk_llm_used),
        "notice": notice,
    }
    if cache is not None and cache_key is not None:
        cache[cache_key] = result
    return result
